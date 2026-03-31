# server.py — AI Career Accelerator Kenya
# FastAPI backend replacing Streamlit
# Deploy on Render (free tier) or any WSGI/ASGI host

from __future__ import annotations
import os
import json
import hashlib
import logging
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Optional

from fastapi import FastAPI, HTTPException, Header, Request
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
import uvicorn

# ── internal modules ──────────────────────────────────────────
from agents import (
    analyze_job,
    rewrite_resume,
    generate_cover_letter,
    generate_emails,
    generate_interview_prep,
    generate_linkedin_optimization,
    generate_gap_analysis,
)
from payments_db import (
    init_db,
    get_or_create_user,
    find_user,
    get_user_tier,
    save_lead,
    mark_user_paid,
    save_user_output,
    load_user_output,
    get_dashboard_stats,
    get_all_leads,
    get_user_payments,
    normalize_phone,
    normalize_email,
    TIER_PRICES,
    TIER_ORDER,
)

# ── optional IntaSend STK push ────────────────────────────────
try:
    from proxy.intasend_client import (
        initiate_stk_push,
        poll_payment_status,
        is_configured as intasend_configured,
        STATE_COMPLETE,
        TERMINAL_STATES,
    )
    _INTASEND = True
except ImportError:
    _INTASEND = False
    intasend_configured = lambda: False

# ── document generation (docx / pdf) ─────────────────────────
try:
    from docx import Document as DocxDocument
    _DOCX = True
except ImportError:
    _DOCX = False

try:
    from fpdf import FPDF
    _PDF = True
except ImportError:
    _PDF = False

# ─────────────────────────────────────────────────────────────
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

init_db()

app = FastAPI(title="AI Career Accelerator Kenya", version="2.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

STATIC_DIR = Path(__file__).parent / "static"
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "")
WHATSAPP_NUMBER = os.getenv("WHATSAPP_NUMBER", "254722285538").strip()


# ─────────────────────────────────────────────────────────────
# AUTH HELPERS
# ─────────────────────────────────────────────────────────────

def _user_id_from_token(token: str) -> str:
    """Token IS the user_id (SHA256 of phone|email, first 24 chars)."""
    return token.strip()


def _require_user(x_user_token: Optional[str] = Header(default=None)) -> str:
    if not x_user_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return _user_id_from_token(x_user_token)


def _require_admin(x_admin_password: Optional[str] = Header(default=None)):
    if not ADMIN_PASSWORD or x_admin_password != ADMIN_PASSWORD:
        raise HTTPException(status_code=403, detail="Forbidden")


def _tier_ok(required: str, user_tier: str) -> bool:
    return TIER_ORDER.index(user_tier) >= TIER_ORDER.index(required)


def _whatsapp_link(user_id: str, phone: str, email: str, tier: str) -> str:
    price = TIER_PRICES.get(tier, 1999)
    tier_name = tier.capitalize()
    msg = (
        f"Hi, I want to pay KES {price:,} for AI Career Accelerator ({tier_name} Plan).\n"
        f"User ID: {user_id}\nPhone: {phone}\nEmail: {email}\n"
        "Please send M-Pesa payment instructions."
    )
    enc = msg.replace("\n", "%0A").replace(" ", "%20")
    return f"https://wa.me/{WHATSAPP_NUMBER}?text={enc}"


# ─────────────────────────────────────────────────────────────
# STATIC FILES
# ─────────────────────────────────────────────────────────────

if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


@app.get("/")
async def root():
    index = STATIC_DIR / "app.html"
    if index.exists():
        return FileResponse(str(index))
    return JSONResponse({"status": "AI Career Accelerator API running"})


@app.get("/admin")
async def admin_page():
    page = STATIC_DIR / "admin.html"
    if page.exists():
        return FileResponse(str(page))
    raise HTTPException(status_code=404, detail="Admin page not found")


# ─────────────────────────────────────────────────────────────
# AUTH
# ─────────────────────────────────────────────────────────────

class LoginRequest(BaseModel):
    phone: str
    email: str


@app.post("/api/login")
async def login(req: LoginRequest):
    user = get_or_create_user(req.phone, req.email)
    if not user:
        raise HTTPException(status_code=400, detail="Invalid phone or email")
    save_lead(req.phone, req.email)
    tier = get_user_tier(user["user_id"])
    return {
        "user_id": user["user_id"],
        "phone": user["phone"],
        "email": user["email"],
        "tier": tier,
        "token": user["user_id"],  # token = user_id
    }


@app.get("/api/me")
async def get_me(x_user_token: Optional[str] = Header(default=None)):
    if not x_user_token:
        raise HTTPException(status_code=401, detail="Not authenticated")
    uid = _user_id_from_token(x_user_token)
    tier = get_user_tier(uid)
    return {"user_id": uid, "tier": tier}


# ─────────────────────────────────────────────────────────────
# JOB ANALYSIS
# ─────────────────────────────────────────────────────────────

class AnalyzeRequest(BaseModel):
    jd: str = Field(..., min_length=30)


@app.post("/api/analyze")
async def api_analyze(req: AnalyzeRequest, x_user_token: Optional[str] = Header(default=None)):
    _require_user(x_user_token)
    result = analyze_job(req.jd)
    return result


# ─────────────────────────────────────────────────────────────
# GENERATE JOB KIT
# ─────────────────────────────────────────────────────────────

class GenerateRequest(BaseModel):
    resume_text: str = Field(..., min_length=20)
    job_analysis: Dict[str, Any]


@app.post("/api/generate")
async def api_generate(req: GenerateRequest, x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    tier = get_user_tier(uid)

    cv = rewrite_resume(req.resume_text, req.job_analysis)
    cover = generate_cover_letter(req.resume_text, req.job_analysis)
    emails = generate_emails(req.job_analysis)
    gap = generate_gap_analysis(req.resume_text, req.job_analysis)

    # Save to DB so user can return later
    save_user_output(uid, cv, cover, emails)

    return {
        "cv": cv,
        "cover_letter": cover,
        "emails": emails,
        "gap": gap,
        "tier": tier,
        "paid": _tier_ok("starter", tier),
    }


@app.post("/api/generate/interview")
async def api_interview(req: GenerateRequest, x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    tier = get_user_tier(uid)
    if not _tier_ok("pro", tier):
        raise HTTPException(status_code=402, detail="Pro plan required")
    result = generate_interview_prep(req.resume_text, req.job_analysis)
    return result


@app.post("/api/generate/linkedin")
async def api_linkedin(req: GenerateRequest, x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    tier = get_user_tier(uid)
    if not _tier_ok("pro", tier):
        raise HTTPException(status_code=402, detail="Pro plan required")
    result = generate_linkedin_optimization(req.resume_text, req.job_analysis)
    return result


# ─────────────────────────────────────────────────────────────
# SAVED OUTPUTS
# ─────────────────────────────────────────────────────────────

@app.get("/api/outputs")
async def get_outputs(x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    saved = load_user_output(uid)
    if not saved:
        return {}
    tier = get_user_tier(uid)
    return {**saved, "tier": tier}


# ─────────────────────────────────────────────────────────────
# DOWNLOADS
# ─────────────────────────────────────────────────────────────

def _md_to_docx(text: str) -> bytes:
    doc = DocxDocument()
    for line in text.splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph("")
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif s.startswith("| "):
            doc.add_paragraph(s)
        else:
            doc.add_paragraph(s)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _md_to_pdf(text: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    for line in text.splitlines():
        s = line.strip()
        if not s:
            pdf.ln(4)
        elif s.startswith("## "):
            pdf.set_font("Arial", "B", 13)
            pdf.multi_cell(0, 7, s[3:])
            pdf.set_font("Arial", size=11)
        elif s.startswith("# "):
            pdf.set_font("Arial", "B", 16)
            pdf.multi_cell(0, 8, s[2:])
            pdf.set_font("Arial", size=11)
        else:
            pdf.multi_cell(0, 6, s)
    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.getvalue()


class DownloadRequest(BaseModel):
    content: str
    format: str  # "md", "docx", "pdf"
    filename: str = "document"


@app.post("/api/download")
async def api_download(req: DownloadRequest, x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    tier = get_user_tier(uid)
    if not _tier_ok("starter", tier):
        raise HTTPException(status_code=402, detail="Starter plan required to download")

    fmt = req.format.lower()
    name = req.filename.replace(" ", "_")

    if fmt == "md":
        return StreamingResponse(
            BytesIO(req.content.encode("utf-8")),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{name}.md"'},
        )
    elif fmt == "docx" and _DOCX:
        data = _md_to_docx(req.content)
        return StreamingResponse(
            BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{name}.docx"'},
        )
    elif fmt == "pdf" and _PDF:
        data = _md_to_pdf(req.content)
        return StreamingResponse(
            BytesIO(data),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{name}.pdf"'},
        )
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")


# ─────────────────────────────────────────────────────────────
# PAYMENT
# ─────────────────────────────────────────────────────────────

class PaymentRequest(BaseModel):
    tier: str
    phone: Optional[str] = None


@app.post("/api/pay/whatsapp")
async def pay_whatsapp(req: PaymentRequest, x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    tier = get_user_tier(uid)
    # Return WhatsApp link — we don't know user's phone/email here without a lookup
    # The client sends phone so we can build the link
    phone = req.phone or ""
    link = _whatsapp_link(uid, phone, "", req.tier)
    return {"whatsapp_url": link, "tier": req.tier, "price": TIER_PRICES.get(req.tier, 1999)}


@app.post("/api/pay/stk")
async def pay_stk(req: PaymentRequest, x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    if not _INTASEND or not intasend_configured():
        raise HTTPException(status_code=503, detail="Automated payment not configured")

    tier = req.tier
    if tier not in TIER_PRICES:
        raise HTTPException(status_code=400, detail="Invalid tier")

    amount = TIER_PRICES[tier]
    result = initiate_stk_push(phone=req.phone, amount=amount, tier=tier, user_id=uid)
    if not result["success"]:
        raise HTTPException(status_code=502, detail=result.get("message", "Payment failed"))

    from payments_db import create_pending_invoice
    create_pending_invoice(result["invoice_id"], uid, req.phone, tier, amount)
    return {"invoice_id": result["invoice_id"], "message": "STK push sent. Check your phone."}


@app.get("/api/pay/poll/{invoice_id}")
async def pay_poll(invoice_id: str, x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    if not _INTASEND:
        raise HTTPException(status_code=503, detail="Payment integration not configured")

    result = poll_payment_status(invoice_id)
    if result.get("paid"):
        from payments_db import confirm_invoice_payment
        confirm_invoice_payment(invoice_id)
        tier = get_user_tier(uid)
        return {"paid": True, "tier": tier}
    return {"paid": False, "state": result.get("state"), "message": result.get("message")}


@app.get("/api/tier")
async def check_tier(x_user_token: Optional[str] = Header(default=None)):
    uid = _require_user(x_user_token)
    tier = get_user_tier(uid)
    return {"tier": tier, "paid": tier != "none"}


# ─────────────────────────────────────────────────────────────
# ADMIN API
# ─────────────────────────────────────────────────────────────

@app.get("/api/admin/stats")
async def admin_stats(x_admin_password: Optional[str] = Header(default=None)):
    _require_admin(x_admin_password)
    return get_dashboard_stats()


@app.get("/api/admin/leads")
async def admin_leads(limit: int = 200, x_admin_password: Optional[str] = Header(default=None)):
    _require_admin(x_admin_password)
    return get_all_leads(limit=limit)


@app.get("/api/admin/payments")
async def admin_payments(limit: int = 100, x_admin_password: Optional[str] = Header(default=None)):
    _require_admin(x_admin_password)
    return get_user_payments(limit=limit)


class AdminUserRequest(BaseModel):
    phone: str
    email: str


@app.post("/api/admin/user")
async def admin_find_user(req: AdminUserRequest, x_admin_password: Optional[str] = Header(default=None)):
    _require_admin(x_admin_password)
    # Use find_user (lookup-only) — never accidentally creates a record
    user = find_user(req.phone, req.email)
    if not user:
        raise HTTPException(status_code=404, detail="User not found. Check phone and email are correct.")
    tier = get_user_tier(user["user_id"])
    payments = get_user_payments(user_id=user["user_id"])
    return {**user, "tier": tier, "payments": payments}


class AdminGrantRequest(BaseModel):
    user_id: str
    tier: str


@app.post("/api/admin/grant")
async def admin_grant(req: AdminGrantRequest, x_admin_password: Optional[str] = Header(default=None)):
    _require_admin(x_admin_password)
    if req.tier not in TIER_PRICES:
        raise HTTPException(status_code=400, detail="Invalid tier")
    ok = mark_user_paid(req.user_id, tier=req.tier)
    if not ok:
        raise HTTPException(status_code=500, detail="Failed to grant tier")
    return {"success": True, "user_id": req.user_id, "tier": req.tier}


# ─────────────────────────────────────────────────────────────
# HEALTH CHECK
# ─────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "intasend": intasend_configured() if _INTASEND else False,
        "docx": _DOCX,
        "pdf": _PDF,
    }


# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    port = int(os.getenv("PORT", 8000))
    uvicorn.run("server:app", host="0.0.0.0", port=port, reload=False)
