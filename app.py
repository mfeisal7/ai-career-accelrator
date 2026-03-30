# app.py — AI Career Accelerator (Rebuilt)
from io import BytesIO
import os
import streamlit as st
from docx import Document
from fpdf import FPDF

from agents import (
    extract_text_from_pdf,
    analyze_job,
    rewrite_resume,
    generate_cover_letter,
    generate_emails,
    generate_linkedin_optimization,
    generate_interview_prep,
    generate_gap_analysis,
)

from payments_db import (
    init_db,
    get_user_payment_status,
    get_user_tier,
    save_user_output,
    load_user_output,
    get_or_create_user,
    save_lead,
    create_pending_invoice,
    confirm_invoice_payment,
    update_invoice_state,
    get_pending_invoice_for_user,
    TIER_PRICES,
)

# IntaSend M-Pesa integration (optional — gracefully degrades to WhatsApp if not configured)
try:
    from proxy.intasend_client import (
        initiate_stk_push,
        poll_payment_status,
        is_configured as intasend_configured,
        STATE_COMPLETE,
        STATE_FAILED,
        STATE_CANCELLED,
        TERMINAL_STATES,
    )
    _INTASEND_AVAILABLE = True
except ImportError:
    _INTASEND_AVAILABLE = False
    intasend_configured = lambda: False

try:
    from streamlit_autorefresh import st_autorefresh
except Exception:
    st_autorefresh = None

# ─────────────────────────────────────────────
# TIER DEFINITIONS
# ─────────────────────────────────────────────
TIERS = {
    "starter": {
        "name": "Starter",
        "price_ksh": 1999,
        "features": [
            "AI-optimised CV (ATS-ready)",
            "Tailored cover letter",
            "4-email follow-up strategy",
            "Job description analysis",
            "Gap analysis report",
        ],
    },
    "pro": {
        "name": "Pro",
        "price_ksh": 3999,
        "features": [
            "Everything in Starter",
            "LinkedIn profile optimization",
            "Full interview prep (13 Q&A)",
            "Salary negotiation script",
            "Priority WhatsApp support",
        ],
    },
    "executive": {
        "name": "Executive",
        "price_ksh": 7999,
        "features": [
            "Everything in Pro",
            "Human review of your CV & cover letter",
            "1-on-1 WhatsApp coaching (30 min)",
            "90-day interview guarantee",
            "Unlimited regenerations",
        ],
    },
}

st.set_page_config(
    page_title="AI Career Accelerator Kenya",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
    .stApp { background-color: #0f172a; color: #f1f5f9; }
    .stTabs [data-baseweb="tab-list"] { gap: 8px; background: #1e293b; border-radius: 12px; padding: 4px; }
    .stTabs [data-baseweb="tab"] { border-radius: 8px; color: #94a3b8; font-weight: 500; }
    .stTabs [aria-selected="true"] { background: #10b981 !important; color: #0f172a !important; }
    div[data-testid="stExpander"] { background: #1e293b; border: 1px solid #334155; border-radius: 12px; }
    .stDownloadButton button { border-radius: 8px; font-weight: 600; background: #059669; color: white; border: none; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def _get_whatsapp_number() -> str:
    return (os.getenv("WHATSAPP_NUMBER") or "254722285538").strip()


def _whatsapp_pay_link(user_id: str, phone: str, email: str, tier: str = "starter") -> str:
    price = TIERS.get(tier, TIERS["starter"])["price_ksh"]
    tier_name = TIERS.get(tier, TIERS["starter"])["name"]
    msg = (
        f"Hi, I want to pay KES {price} for AI Career Accelerator ({tier_name} Plan).\n"
        f"User ID: {user_id}\nPhone: {phone}\nEmail: {email}\n"
        "Please send M-Pesa payment instructions."
    )
    msg_encoded = msg.replace("\n", "%0A").replace(" ", "%20")
    return f"https://wa.me/{_get_whatsapp_number()}?text={msg_encoded}"


def _hydrate(user_id: str):
    saved = load_user_output(user_id)
    if not saved:
        return
    for key, val in saved.items():
        st.session_state.setdefault(key, val)


def _markdown_to_docx(text: str) -> bytes:
    doc = Document()
    for line in text.splitlines():
        s = line.strip()
        if not s:
            doc.add_paragraph("")
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- ") or s.startswith("* "):
            doc.add_paragraph(s[2:], style="List Bullet")
        else:
            doc.add_paragraph(s)
    buff = BytesIO()
    doc.save(buff)
    buff.seek(0)
    return buff.getvalue()


def _markdown_to_pdf(text: str) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=11)
    for line in text.splitlines():
        s = line.strip()
        if not s:
            pdf.ln(4)
        elif s.startswith("## "):
            pdf.set_font("Arial", "B", 13)
            pdf.multi_cell(0, 7, s[3:])
            pdf.set_font("Arial", size=11)
        elif s.startswith("# "):
            pdf.set_font("Arial", "B", 16)
            pdf.multi_cell(0, 8, s[2:])
            pdf.set_font("Arial", size=11)
        else:
            pdf.multi_cell(0, 6, s)
    buff = BytesIO()
    pdf.output(buff)
    buff.seek(0)
    return buff.getvalue()


def _tier_check(required_tier: str, user_tier: str) -> bool:
    order = ["none", "starter", "pro", "executive"]
    return order.index(user_tier) >= order.index(required_tier)


# ─────────────────────────────────────────────
# LOGIN
# ─────────────────────────────────────────────

def _require_login():
    if st.session_state.get("user_id"):
        return

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.markdown("""
        <div style='text-align:center; padding: 40px 0 20px 0;'>
            <div style='font-size:48px;'>🚀</div>
            <h1 style='color:#10b981; margin:8px 0 4px 0;'>AI Career Accelerator</h1>
            <p style='color:#94a3b8;'>Kenya's AI-powered job kit. Built for Kenyan graduates.</p>
        </div>
        """, unsafe_allow_html=True)

        st.info("Enter your phone number and email to access the app. New users get a free account instantly.")

        with st.form("login_form", clear_on_submit=False):
            phone = st.text_input("📱 Phone Number", placeholder="0722 123 456 or +254722123456")
            email = st.text_input("📧 Email Address", placeholder="you@gmail.com")
            submitted = st.form_submit_button("→ Enter App", use_container_width=True)

        if not submitted:
            st.stop()

        phone = (phone or "").strip()
        email = (email or "").strip()

        if not phone or not email:
            st.error("Please enter both phone number and email.")
            st.stop()

        with st.spinner("Setting up your account…"):
            user = get_or_create_user(phone, email)
            save_lead(phone, email)

        if not user:
            st.error("Invalid phone or email. Please try again.")
            st.stop()

        st.session_state["user_id"] = user["user_id"]
        st.session_state["user_phone"] = user["phone"]
        st.session_state["user_email"] = user["email"]
        st.rerun()


# ─────────────────────────────────────────────
# M-PESA PAYMENT GATE (automated STK push + WhatsApp fallback)
# ─────────────────────────────────────────────

def _payment_gate(feature_name: str, required_tier: str, user_id: str, user_tier: str):
    """
    Full payment gate with:
    - Automated M-Pesa STK push (if IntaSend is configured)
    - Live polling (no page refresh needed)
    - WhatsApp manual fallback
    - Manual "I've paid" check
    """
    tier_info = TIERS.get(required_tier, TIERS["starter"])
    phone = st.session_state.get("user_phone", "")
    email = st.session_state.get("user_email", "")
    price = tier_info["price_ksh"]
    gate_key = f"gate_{required_tier}_{feature_name}"

    # ── Locked header ──
    st.markdown(f"""
    <div style='background:#1e293b; border:1px solid #334155; border-radius:12px;
                padding:20px 24px; margin-bottom:16px;'>
        <div style='display:flex; align-items:center; gap:12px;'>
            <span style='font-size:28px;'>🔒</span>
            <div>
                <div style='font-weight:700; color:#f1f5f9; font-size:15px;'>{feature_name}</div>
                <div style='color:#94a3b8; font-size:13px;'>
                    Requires the <strong style='color:#10b981;'>{tier_info["name"]} Plan</strong>
                    — KSh {price:,} (one-time, lifetime access)
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    use_automated = _INTASEND_AVAILABLE and intasend_configured()

    # ── Check if there's already a pending invoice ──
    pending = get_pending_invoice_for_user(user_id)
    if pending and pending.get("tier") == required_tier:
        invoice_id = pending["invoice_id"]
        st.info("⏳ Payment in progress — waiting for M-Pesa confirmation…")

        col_poll, col_cancel = st.columns([3, 1])
        with col_poll:
            if st.button("🔄 Check Payment Status", key=f"poll_{gate_key}", use_container_width=True):
                result = poll_payment_status(invoice_id)
                if result["paid"]:
                    confirm_invoice_payment(invoice_id)
                    st.session_state["user_tier"] = get_user_tier(user_id)
                    st.success("✅ Payment confirmed! Access unlocked.")
                    st.rerun()
                elif result["state"] in TERMINAL_STATES:
                    update_invoice_state(invoice_id, result["state"])
                    st.error(result["message"])
                    st.session_state.pop(f"pending_invoice_{gate_key}", None)
                    st.rerun()
                else:
                    st.info(result["message"])
        with col_cancel:
            if st.button("✕ Cancel", key=f"cancel_{gate_key}", use_container_width=True):
                update_invoice_state(invoice_id, "CANCELLED")
                st.rerun()

        # Auto-refresh to keep polling
        if st_autorefresh is not None:
            st_autorefresh(interval=5000, key=f"refresh_{gate_key}")
        return

    # ── Payment options ──
    if use_automated:
        st.markdown(f"**Pay KSh {price:,} via M-Pesa** — enter your phone and we'll send a payment prompt directly to it.")

        pay_phone = st.text_input(
            "📱 M-Pesa phone number",
            value=phone,
            key=f"pay_phone_{gate_key}",
            placeholder="0722 123 456",
        )

        col_pay, col_wa = st.columns([2, 1])
        with col_pay:
            if st.button(
                f"🏦 Send M-Pesa Prompt — KSh {price:,}",
                key=f"pay_{gate_key}",
                use_container_width=True,
                type="primary",
            ):
                if not pay_phone or not pay_phone.strip():
                    st.error("Enter your M-Pesa phone number.")
                else:
                    with st.spinner("Sending M-Pesa payment prompt to your phone…"):
                        result = initiate_stk_push(
                            phone=pay_phone.strip(),
                            amount=price,
                            tier=required_tier,
                            user_id=user_id,
                        )

                    if result["success"]:
                        invoice_id = result["invoice_id"]
                        create_pending_invoice(invoice_id, user_id, pay_phone.strip(), required_tier, price)
                        st.success(f"✅ M-Pesa prompt sent! Check your phone and enter your PIN.")
                        st.info("This page will check every 5 seconds. You can also click 'Check Payment Status' below.")
                        st.rerun()
                    else:
                        st.error(f"Could not initiate payment: {result['message']}")
                        st.info("Use the WhatsApp option below instead.")

        with col_wa:
            wa_link = _whatsapp_pay_link(user_id, phone, email, required_tier)
            st.link_button("💬 WhatsApp instead", wa_link, use_container_width=True)

    else:
        # WhatsApp-only mode (IntaSend not configured)
        st.markdown(f"""
        <div style='background:#0d2d1f; border:1px solid #10b981; border-radius:10px; padding:16px; margin-bottom:12px;'>
            <div style='font-weight:600; color:#10b981; margin-bottom:6px;'>Pay via M-Pesa</div>
            <ol style='color:#cbd5e1; font-size:13px; margin:0; padding-left:18px; line-height:1.8;'>
                <li>Click "Message us on WhatsApp" below</li>
                <li>We'll send you M-Pesa Paybill instructions</li>
                <li>Pay via M-Pesa (30 seconds)</li>
                <li>Come back and click "Check my access" — usually unlocks within 15 min</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)

        wa_link = _whatsapp_pay_link(user_id, phone, email, required_tier)
        st.link_button(
            f"💬 Message us on WhatsApp — KSh {price:,}",
            wa_link,
            use_container_width=True,
        )

    st.markdown("")
    if st.button(f"✅ I've already paid — check my access", key=f"check_{gate_key}"):
        new_tier = get_user_tier(user_id)
        if _tier_check(required_tier, new_tier):
            st.session_state["user_tier"] = new_tier
            st.success("✅ Access confirmed!")
            st.rerun()
        else:
            st.warning("Payment not confirmed yet. If you've paid via WhatsApp, please wait for confirmation (usually within 15 min) or contact us.")


# Keep _locked_feature as alias for backwards compat
_locked_feature = _payment_gate


# ─────────────────────────────────────────────
# DOWNLOAD BLOCK
# ─────────────────────────────────────────────

def _download_block(title: str, content: str, filename_stem: str):
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button(
            "⬇️ Download .md",
            content.encode("utf-8"),
            f"{filename_stem}.md",
            "text/markdown",
            use_container_width=True,
        )
    with c2:
        st.download_button(
            "⬇️ Download .docx",
            _markdown_to_docx(content),
            f"{filename_stem}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with c3:
        st.download_button(
            "⬇️ Download .pdf",
            _markdown_to_pdf(content),
            f"{filename_stem}.pdf",
            "application/pdf",
            use_container_width=True,
        )


# ─────────────────────────────────────────────
# MAIN APP
# ─────────────────────────────────────────────

def main():
    init_db()
    _require_login()

    user_id = st.session_state["user_id"]
    user_phone = st.session_state.get("user_phone", "")
    user_email = st.session_state.get("user_email", "")

    if "user_tier" not in st.session_state:
        st.session_state["user_tier"] = get_user_tier(user_id)
    user_tier = st.session_state["user_tier"]

    _hydrate(user_id)

    # Header
    col_logo, col_user = st.columns([3, 1])
    with col_logo:
        st.markdown("""
        <h2 style='margin:0; color:#10b981;'>🚀 AI Career Accelerator</h2>
        <p style='margin:0; color:#64748b; font-size:13px;'>Built for Kenyan graduates & early-career professionals</p>
        """, unsafe_allow_html=True)
    with col_user:
        tier_label = user_tier.upper() if user_tier != "none" else "FREE"
        st.markdown(f"""
        <div style='text-align:right; padding-top:4px;'>
            <span style='color:#94a3b8; font-size:12px;'>{user_phone}</span><br>
            <span style='background:#059669; color:white; border-radius:20px; padding:2px 10px; font-size:12px; font-weight:600;'>{tier_label}</span>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    tab_job, tab_resume, tab_downloads, tab_linkedin, tab_interview, tab_pricing = st.tabs([
        "1️⃣ Job Analysis",
        "2️⃣ My CV",
        "3️⃣ Downloads",
        "4️⃣ LinkedIn",
        "5️⃣ Interview Prep",
        "💳 Plans",
    ])

    # ── TAB 1: JOB ANALYSIS ──
    with tab_job:
        st.subheader("Step 1: Paste the job you're applying for")
        st.caption("The AI extracts keywords, requirements, and insights to tailor everything.")

        jd = st.text_area(
            "Job Description",
            height=280,
            placeholder="Paste the full job description here…",
            key="jd_input",
        )

        if st.button("🔍 Analyse Job Description", use_container_width=True, type="primary"):
            if not jd.strip():
                st.error("Please paste a job description first.")
            else:
                with st.spinner("Analysing…"):
                    analysis = analyze_job(jd)
                st.session_state["job_analysis"] = analysis
                st.session_state["job_raw_jd"] = jd
                st.success("✅ Analysed! See the breakdown below.")

        if "job_analysis" in st.session_state:
            a = st.session_state["job_analysis"]

            col1, col2, col3, col4 = st.columns(4)
            col1.metric("Role", a.get("job_title", "—"))
            col2.metric("Sector", a.get("sector", "—"))
            col3.metric("Seniority", a.get("seniority", "—"))
            col4.metric("Est. Salary", a.get("salary_estimate", "—"))

            with st.expander("📋 Key Requirements"):
                for r in a.get("key_requirements", []):
                    st.markdown(f"- {r}")

            with st.expander("🔑 ATS Keywords to use in your CV"):
                st.markdown(" • ".join(f"`{k}`" for k in a.get("ats_keywords", [])))

            with st.expander("💡 Application Tips"):
                for tip in a.get("application_tips", []):
                    st.markdown(f"- {tip}")

            if a.get("red_flags"):
                with st.expander("⚠️ Red Flags / Clarify at Interview"):
                    for flag in a.get("red_flags", []):
                        st.markdown(f"- {flag}")

    # ── TAB 2: MY CV ──
    with tab_resume:
        st.subheader("Step 2: Generate your AI-powered CV & Cover Letter")

        if "job_analysis" not in st.session_state:
            st.warning("⬆️ Analyse a job in Tab 1 first.")
        else:
            mode = st.radio(
                "How to provide your background?",
                ["📝 Fill in Profile Form", "📄 Upload PDF", "✏️ Paste Text"],
                horizontal=True,
                help="The Profile Form gives the best results — each section is clearly structured.",
            )

            resume_text = ""

            # ── STRUCTURED PROFILE FORM ──
            if mode == "📝 Fill in Profile Form":
                st.info("Fill in each section below. The more detail you add, the stronger your CV will be.")

                with st.form("profile_builder_form", clear_on_submit=False):
                    st.markdown("#### 👤 Personal Details")
                    col_n, col_p2, col_e2 = st.columns(3)
                    with col_n:
                        p_name = st.text_input("Full name*", placeholder="Jane Wanjiku")
                    with col_p2:
                        p_phone = st.text_input("Phone*", placeholder="0712 345 678")
                    with col_e2:
                        p_email = st.text_input("Email*", placeholder="jane@gmail.com")
                    col_loc, col_li = st.columns(2)
                    with col_loc:
                        p_location = st.text_input("Location", placeholder="Nairobi, Kenya")
                    with col_li:
                        p_linkedin = st.text_input("LinkedIn URL (optional)", placeholder="linkedin.com/in/janewanjiku")

                    st.markdown("---")
                    st.markdown("#### 🎓 Education")
                    st.caption("Add your highest qualifications. Use one entry per line.")

                    edu_entries = []
                    for i in range(1, 4):
                        with st.expander(f"Qualification {i}" + (" (required)" if i == 1 else " (optional)"), expanded=(i == 1)):
                            col_d, col_i, col_y = st.columns([2, 2, 1])
                            with col_d:
                                degree = st.text_input(f"Degree/Diploma", placeholder="BSc Computer Science", key=f"edu_degree_{i}")
                            with col_i:
                                institution = st.text_input(f"Institution", placeholder="University of Nairobi", key=f"edu_inst_{i}")
                            with col_y:
                                year = st.text_input(f"Year", placeholder="2020", key=f"edu_year_{i}")
                            honours = st.text_input(f"Grade/Honours (optional)", placeholder="Second Class Upper / GPA 3.5", key=f"edu_grade_{i}")
                            if degree and institution:
                                edu_entries.append({"degree": degree, "institution": institution, "year": year, "honours": honours})

                    st.markdown("---")
                    st.markdown("#### 💼 Work Experience")
                    st.caption("Most recent first. Use bullet points in the Responsibilities field — start each with a dash (-) or new line.")

                    work_entries = []
                    for i in range(1, 5):
                        with st.expander(f"Job {i}" + (" (required)" if i == 1 else " (optional)"), expanded=(i == 1)):
                            col_r, col_c2 = st.columns([2, 2])
                            with col_r:
                                role = st.text_input("Job Title", placeholder="Data Analyst", key=f"work_role_{i}")
                            with col_c2:
                                company = st.text_input("Company / Organisation", placeholder="Equity Bank", key=f"work_company_{i}")
                            col_s, col_e_w = st.columns(2)
                            with col_s:
                                start = st.text_input("Start date", placeholder="Jan 2021", key=f"work_start_{i}")
                            with col_e_w:
                                end = st.text_input("End date", placeholder="Present", key=f"work_end_{i}")
                            responsibilities = st.text_area(
                                "Key responsibilities & achievements (one per line)",
                                height=120,
                                placeholder="- Analysed financial data and produced monthly reports for senior management\n- Reduced report production time by 30% by automating Excel models\n- Managed a team of 3 junior analysts",
                                key=f"work_resp_{i}",
                            )
                            if role and company:
                                work_entries.append({
                                    "role": role, "company": company,
                                    "start": start, "end": end,
                                    "responsibilities": responsibilities,
                                })

                    st.markdown("---")
                    st.markdown("#### 🛠️ Skills & Qualifications")
                    col_sk1, col_sk2 = st.columns(2)
                    with col_sk1:
                        p_skills = st.text_area(
                            "Technical skills (comma-separated)",
                            height=80,
                            placeholder="Python, Excel, SQL, Power BI, QuickBooks, SPSS",
                        )
                    with col_sk2:
                        p_soft = st.text_area(
                            "Soft skills (comma-separated)",
                            height=80,
                            placeholder="Team leadership, Communication, Problem solving, Project management",
                        )

                    p_certs = st.text_area(
                        "Certifications & Professional memberships (one per line, optional)",
                        height=70,
                        placeholder="CPA Part II — KASNEB\nGoogle Data Analytics Certificate — Coursera, 2023\nICSK Member",
                    )

                    p_extra = st.text_area(
                        "Volunteer work, Projects, Awards (optional)",
                        height=70,
                        placeholder="Volunteer data trainer — Nairobi Tech Week, 2023\nBest Accounting Student Award — JKUAT, 2020",
                    )

                    build_submitted = st.form_submit_button("🔨 Build Structured Profile", use_container_width=True)

                if build_submitted:
                    if not p_name.strip() or not edu_entries or not work_entries:
                        st.error("Please fill in at least: your name, one education entry, and one work experience entry.")
                    else:
                        # Assemble clean structured text for the template engine
                        lines = []
                        lines.append(p_name.strip())
                        contact_parts = [p for p in [p_phone.strip(), p_email.strip(), p_location.strip()] if p]
                        lines.append(" | ".join(contact_parts))
                        if p_linkedin.strip():
                            lines.append(p_linkedin.strip())
                        lines.append("")

                        lines.append("Education:")
                        for edu in edu_entries:
                            entry_line = f"{edu['degree']} - {edu['institution']}"
                            if edu.get("year"):
                                entry_line += f" - {edu['year']}"
                            if edu.get("honours"):
                                entry_line += f" ({edu['honours']})"
                            lines.append(entry_line)
                        lines.append("")

                        lines.append("Work Experience:")
                        for job in work_entries:
                            date_range = f"{job['start']}" + (f" to {job['end']}" if job.get("end") else "")
                            lines.append(f"{job['role']} - {job['company']} - {date_range}")
                            if job.get("responsibilities"):
                                for resp_line in job["responsibilities"].splitlines():
                                    resp_line = resp_line.strip().lstrip("-• ").strip()
                                    if resp_line:
                                        lines.append(f"- {resp_line}")
                            lines.append("")

                        all_skills = []
                        if p_skills.strip():
                            all_skills.extend([s.strip() for s in p_skills.split(",") if s.strip()])
                        if p_soft.strip():
                            all_skills.extend([s.strip() for s in p_soft.split(",") if s.strip()])
                        if all_skills:
                            lines.append("Skills:")
                            lines.append(", ".join(all_skills))
                            lines.append("")

                        if p_certs.strip():
                            lines.append("Certifications:")
                            for cert_line in p_certs.splitlines():
                                if cert_line.strip():
                                    lines.append(cert_line.strip())
                            lines.append("")

                        if p_extra.strip():
                            lines.append("Volunteer & Achievements:")
                            for extra_line in p_extra.splitlines():
                                if extra_line.strip():
                                    lines.append(extra_line.strip())
                            lines.append("")

                        resume_text = "\n".join(lines)
                        st.session_state["profile_form_text"] = resume_text
                        st.success("✅ Profile built! Ready to generate your job kit.")

                # Use saved profile form text if available
                if not resume_text:
                    resume_text = st.session_state.get("profile_form_text", "")

                if resume_text:
                    with st.expander("📋 Preview your structured profile", expanded=False):
                        st.text(resume_text)

            elif mode == "📄 Upload PDF":
                f = st.file_uploader("Upload current CV (PDF)", type=["pdf"])
                if f is not None:
                    with st.spinner("Extracting text…"):
                        resume_text = extract_text_from_pdf(f)
                    resume_text = st.text_area("Extracted text (editable)", resume_text, height=200, key="pdf_extract")
            else:
                resume_text = st.text_area(
                    "Paste your CV or describe your experience",
                    height=250,
                    placeholder="Include: name, education, work history, skills, projects, volunteer work, achievements…",
                    key="manual_resume",
                )

            if st.button("⚡ Generate My AI Job Kit", use_container_width=True, type="primary"):
                if not (resume_text or "").strip():
                    st.error("Please provide your CV or experience.")
                else:
                    prog = st.progress(0, text="Starting…")
                    prog.progress(15, text="Rewriting your CV with ATS keywords…")
                    ai_resume = rewrite_resume(resume_text, st.session_state["job_analysis"])
                    prog.progress(40, text="Crafting your cover letter…")
                    ai_cover = generate_cover_letter(resume_text, st.session_state["job_analysis"])
                    prog.progress(60, text="Building email strategy…")
                    ai_emails = generate_emails(st.session_state["job_analysis"])
                    prog.progress(80, text="Running gap analysis…")
                    gap = generate_gap_analysis(resume_text, st.session_state["job_analysis"])
                    prog.progress(100, text="Done!")

                    st.session_state["ai_resume_markdown"] = ai_resume
                    st.session_state["ai_cover_letter"] = ai_cover
                    st.session_state["ai_emails"] = ai_emails
                    st.session_state["gap_analysis"] = gap
                    st.session_state["resume_text_raw"] = resume_text

                    save_user_output(user_id, ai_resume, ai_cover, ai_emails)
                    st.success("✅ Your AI Job Kit is ready! Go to the Downloads tab.")

                    if gap:
                        score = gap.get("match_score", 0)
                        color = "#10b981" if score >= 70 else "#f59e0b" if score >= 50 else "#ef4444"
                        st.markdown(f"""
                        <div style='background:#1e293b; border-radius:12px; padding:20px; margin-top:16px;'>
                            <h4 style='margin:0 0 4px 0;'>Gap Analysis</h4>
                            <div style='font-size:40px; font-weight:800; color:{color};'>{score}%</div>
                            <p style='color:#94a3b8; margin:0;'>match with job requirements</p>
                        </div>
                        """, unsafe_allow_html=True)
                        col_s, col_g = st.columns(2)
                        with col_s:
                            st.markdown("**✅ Your Strengths**")
                            for s in gap.get("strengths", []):
                                st.markdown(f"- {s}")
                        with col_g:
                            st.markdown("**⚠️ Gaps to Close**")
                            for g in gap.get("gaps", []):
                                st.markdown(f"- {g}")

    # ── TAB 3: DOWNLOADS ──
    with tab_downloads:
        st.subheader("Your AI Job Kit — Downloads")
        resume = st.session_state.get("ai_resume_markdown", "")
        cover = st.session_state.get("ai_cover_letter", "")
        emails = st.session_state.get("ai_emails", [])

        if not resume and not cover:
            st.info("Generate your job kit in Tab 2 first.")
        else:
            paid = _tier_check("starter", user_tier)

            st.markdown("### 📄 AI-Optimised CV")
            with st.expander("Preview CV", expanded=True):
                st.markdown(resume)
            if paid:
                _download_block("CV", resume, "AI_CV_Kenya")
            else:
                _locked_feature("CV Downloads", "starter", user_id, user_tier)

            st.markdown("---")
            st.markdown("### ✉️ Cover Letter")
            with st.expander("Preview Cover Letter"):
                st.markdown(cover)
            if paid:
                _download_block("Cover Letter", cover, "Cover_Letter_Kenya")
            else:
                _locked_feature("Cover Letter Downloads", "starter", user_id, user_tier)

            st.markdown("---")
            st.markdown("### 📬 4-Email Follow-Up Strategy")
            if emails:
                for i, em in enumerate(emails):
                    if isinstance(em, dict):
                        with st.expander(f"Email {i+1}: {em.get('purpose', '')} — {em.get('timing', '')}"):
                            st.markdown(f"**Subject:** {em.get('subject', '')}")
                            st.markdown("---")
                            st.markdown(em.get("body", ""))

    # ── TAB 4: LINKEDIN (Pro+) ──
    with tab_linkedin:
        st.subheader("LinkedIn Profile Optimization")
        st.caption("Optimized headline, About section, bullet points, and cold outreach templates.")

        if not _tier_check("pro", user_tier):
            _locked_feature("LinkedIn Optimization", "pro", user_id, user_tier)
        elif "job_analysis" not in st.session_state:
            st.warning("⬆️ Analyse a job in Tab 1 first.")
        else:
            resume_raw = st.session_state.get("resume_text_raw", st.session_state.get("ai_resume_markdown", ""))
            if not resume_raw:
                st.warning("⬆️ Generate your CV in Tab 2 first.")
            else:
                if st.button("🔗 Optimize My LinkedIn", use_container_width=True, type="primary"):
                    with st.spinner("Optimizing…"):
                        li = generate_linkedin_optimization(resume_raw, st.session_state["job_analysis"])
                    st.session_state["linkedin_data"] = li

                if "linkedin_data" in st.session_state:
                    li = st.session_state["linkedin_data"]
                    st.markdown("### Headline")
                    st.code(li.get("headline", ""), language=None)
                    st.markdown("### About Section")
                    st.text_area("Copy to LinkedIn", li.get("about", ""), height=200)
                    st.markdown("### Experience Bullet Points")
                    for b in li.get("experience_bullets", []):
                        st.markdown(f"- {b}")
                    col_sk, col_msg = st.columns(2)
                    with col_sk:
                        st.markdown("### Skills to Add")
                        for s in li.get("featured_skills", []):
                            st.markdown(f"- `{s}`")
                    with col_msg:
                        st.markdown("### Connection Request (300 chars)")
                        st.text_area("", li.get("connection_message", ""), height=100)
                    st.markdown("### Cold InMail Template")
                    st.text_area("Send to hiring managers", li.get("inmail_template", ""), height=150)

    # ── TAB 5: INTERVIEW PREP (Pro+) ──
    with tab_interview:
        st.subheader("Interview Preparation Guide")
        st.caption("13 tailored Q&As, salary negotiation script, first impression tips.")

        if not _tier_check("pro", user_tier):
            _locked_feature("Interview Preparation", "pro", user_id, user_tier)
        elif "job_analysis" not in st.session_state:
            st.warning("⬆️ Analyse a job in Tab 1 first.")
        else:
            resume_raw = st.session_state.get("resume_text_raw", st.session_state.get("ai_resume_markdown", ""))
            if not resume_raw:
                st.warning("⬆️ Generate your CV in Tab 2 first.")
            else:
                if st.button("🎤 Generate Interview Guide", use_container_width=True, type="primary"):
                    with st.spinner("Preparing your guide…"):
                        prep = generate_interview_prep(resume_raw, st.session_state["job_analysis"])
                    st.session_state["interview_prep"] = prep

                if "interview_prep" in st.session_state:
                    prep = st.session_state["interview_prep"]

                    for section_key, section_title in [
                        ("common_questions", "💬 Common Questions"),
                        ("technical_questions", "🔧 Technical / Role Questions"),
                        ("behavioural_questions", "🌟 Behavioural Questions (STAR)"),
                    ]:
                        st.markdown(f"### {section_title}")
                        for q in prep.get(section_key, []):
                            if isinstance(q, dict):
                                with st.expander(q.get("question", "Question")):
                                    st.markdown(f"**Answer:** {q.get('strong_answer', '')}")
                                    if q.get("tip"):
                                        st.info(f"💡 {q['tip']}")

                    st.markdown("### ❓ Questions to Ask Them")
                    for q in prep.get("questions_to_ask_employer", []):
                        st.markdown(f"- {q}")

                    st.markdown("### 💰 Salary Negotiation Script")
                    st.info(prep.get("salary_negotiation_script", ""))

                    st.markdown("### ✅ First Impression Tips")
                    for tip in prep.get("first_impression_tips", []):
                        st.markdown(f"- {tip}")

    # ── TAB 6: PRICING ──
    with tab_pricing:
        st.subheader("Plans & Pricing — One-time. Lifetime access.")

        col_s, col_p, col_e = st.columns(3)

        # ── STARTER ──
        with col_s:
            st.markdown("### Starter")
            st.markdown("## KSh 1,999")
            st.caption("One-time • Lifetime access")
            st.divider()
            for f in TIERS["starter"]["features"]:
                st.markdown(f"✅ {f}")
            st.divider()
            wa_link = _whatsapp_pay_link(user_id, user_phone, user_email, "starter")
            st.link_button("💬 Get Starter — KSh 1,999", wa_link, use_container_width=True)

        # ── PRO (most popular) ──
        with col_p:
            st.markdown("### 🔥 Pro — *Most Popular*")
            st.markdown("## KSh 3,999")
            st.caption("One-time • Lifetime access")
            st.divider()
            for f in TIERS["pro"]["features"]:
                st.markdown(f"✅ {f}")
            st.divider()
            wa_link = _whatsapp_pay_link(user_id, user_phone, user_email, "pro")
            st.link_button("💬 Get Pro — KSh 3,999", wa_link, use_container_width=True, type="primary")

        # ── EXECUTIVE ──
        with col_e:
            st.markdown("### Executive")
            st.markdown("## KSh 7,999")
            st.caption("One-time • Lifetime access")
            st.divider()
            for f in TIERS["executive"]["features"]:
                st.markdown(f"✅ {f}")
            st.divider()
            wa_link = _whatsapp_pay_link(user_id, user_phone, user_email, "executive")
            st.link_button("💬 Get Executive — KSh 7,999", wa_link, use_container_width=True)

        st.divider()
        st.info("🛡️ **45-Day Interview Guarantee** — Follow the system, apply to agreed roles, and if you don't get 3 interviews in 45 days we refund you. No drama.")

        st.markdown("**How payment works:**")
        st.markdown(
            "1. Click the WhatsApp button for your plan\n"
            "2. We send you M-Pesa payment instructions\n"
            "3. Pay via M-Pesa (30 seconds)\n"
            "4. We confirm and unlock your plan (within 15 min, often faster)\n"
            "5. Return here and all features unlock automatically"
        )


if __name__ == "__main__":
    main()
